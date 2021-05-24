from docx import Document
import glob
import os
from docx.shared import Inches
from datetime import date

### file path
path = (r"C:\Users\maxwl\PycharmProjects\test")

### file creation
document = Document()

### add header
document.add_heading('Nice Health check', level=1)

### loop to find all *.pngs

for root, dirs, files in os.walk(path):
    for file in files:
            if file.endswith(".png"):
                document.add_picture(file, width=Inches(8))


### -- save with today's date
today = date.today()
d3 = today.strftime("%m.%d.%y")
document.save(f'{d3}_Nice_Health_Check.docx')

